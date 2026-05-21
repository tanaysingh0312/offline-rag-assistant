import fitz  # PyMuPDF
import docx
from langchain_core.documents import Document
import os

def parse_pdf(file_path: str, filename: str) -> list[Document]:
    docs = []
    doc = fitz.open(file_path)
    for page_num, page in enumerate(doc):
        text = page.get_text()
        if text.strip():
            docs.append(
                Document(
                    page_content=text,
                    metadata={"filename": filename, "page_number": page_num + 1}
                )
            )
    return docs

def parse_docx(file_path: str, filename: str) -> list[Document]:
    doc = docx.Document(file_path)
    text = "\n".join([para.text for para in doc.paragraphs])
    return [Document(page_content=text, metadata={"filename": filename, "page_number": 1})]

def parse_txt(file_path: str, filename: str) -> list[Document]:
    with open(file_path, "r", encoding="utf-8") as f:
        text = f.read()
    return [Document(page_content=text, metadata={"filename": filename, "page_number": 1})]

def parse_document(file_path: str, filename: str) -> list[Document]:
    ext = os.path.splitext(filename)[1].lower()
    if ext == ".pdf":
        return parse_pdf(file_path, filename)
    elif ext == ".docx":
        return parse_docx(file_path, filename)
    elif ext == ".txt":
        return parse_txt(file_path, filename)
    else:
        raise ValueError(f"Unsupported file extension: {ext}")
